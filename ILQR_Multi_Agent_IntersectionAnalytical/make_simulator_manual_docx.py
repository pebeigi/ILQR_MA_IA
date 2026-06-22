"""Build a complete Word manual for the ILQ multi-agent simulator."""
from __future__ import annotations

from pathlib import Path
from typing import Iterable

import numpy as np
from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

import main


OUT_DIR = Path(__file__).resolve().parent / "outputs" / "docs"
OUT_PATH = OUT_DIR / "ILQ_Multi_Agent_Intersection_Simulator_Manual.docx"

BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
INK = RGBColor(0x22, 0x22, 0x22)
MUTED = RGBColor(0x66, 0x66, 0x66)
TABLE_HEADER_FILL = "E8EEF5"
CALLOUT_FILL = "F4F6F9"
CODE_FILL = "F2F4F7"
BORDER = "B8C6D8"

CONTENT_DXA = 9360
TABLE_INDENT_DXA = 120
CELL_MARGINS = {"top": 80, "bottom": 80, "start": 120, "end": 120}


def set_run_font(run, *, name="Calibri", size=None, bold=None, italic=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, margins=None):
    margins = CELL_MARGINS if margins is None else margins
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, value in margins.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_dxa)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: list[int], *, indent_dxa: int = TABLE_INDENT_DXA):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(int(width)))
        grid.append(grid_col)
    tbl.insert(0, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def paragraph_border(paragraph, *, bottom=True, color=BORDER, size="8", space="4"):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    tag = "bottom" if bottom else "top"
    node = p_bdr.find(qn(f"w:{tag}"))
    if node is None:
        node = OxmlElement(f"w:{tag}")
        p_bdr.append(node)
    node.set(qn("w:val"), "single")
    node.set(qn("w:sz"), size)
    node.set(qn("w:space"), space)
    node.set(qn("w:color"), color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)
    set_run_font(run, size=9, color=MUTED)


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "ILQ Multi-Agent Intersection Simulator Manual"
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(header.runs[0], size=9, color=MUTED)
    paragraph_border(header, bottom=True, color="D9E2EF", size="4", space="2")

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer.add_run("Page ")
    set_run_font(r, size=9, color=MUTED)
    add_page_number(footer)


def add_paragraph(doc, text: str, *, style=None, bold_prefix: str | None = None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        rb = p.add_run(bold_prefix)
        set_run_font(rb, bold=True)
        rest = text[len(bold_prefix):]
        if rest:
            r = p.add_run(rest)
            set_run_font(r)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_bullet(doc, text: str):
    return add_paragraph(doc, text, style="List Bullet")


def add_numbered(doc, text: str):
    return add_paragraph(doc, text, style="List Number")


def add_ordered_list(doc, items: Iterable[str]):
    for idx, text in enumerate(items, start=1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.375)
        p.paragraph_format.first_line_indent = Inches(-0.188)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.25
        r_num = p.add_run(f"{idx}.  ")
        set_run_font(r_num)
        r_text = p.add_run(text)
        set_run_font(r_text)


def add_formula(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.05
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), CODE_FILL)
    p_pr.append(shd)
    run = p.add_run(text)
    set_run_font(run, name="Courier New", size=9.2, color=RGBColor(0x20, 0x20, 0x20))
    return p


def add_callout(doc, title: str, body: str):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_DXA - TABLE_INDENT_DXA], indent_dxa=TABLE_INDENT_DXA)
    cell = table.cell(0, 0)
    set_cell_shading(cell, CALLOUT_FILL)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5)
    return table


def add_table(doc, headers: list[str], rows: Iterable[Iterable[object]], widths_in: list[float]):
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    widths_dxa = [round(w * 1440) for w in widths_in]
    total = sum(widths_dxa)
    if total != CONTENT_DXA:
        widths_dxa[-1] += CONTENT_DXA - total
    set_table_geometry(table, widths_dxa)

    for c, header in enumerate(headers):
        cell = table.cell(0, c)
        set_cell_shading(cell, TABLE_HEADER_FILL)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(header))
        set_run_font(r, size=9.5, bold=True, color=DARK_BLUE)

    for r_idx, row in enumerate(rows, start=1):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            if r_idx % 2 == 0:
                set_cell_shading(cell, "FAFBFC")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run("" if value is None else str(value))
            set_run_font(run, size=9.2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def fmt_array(arr):
    vals = np.asarray(arr, dtype=float).reshape(-1)
    return "[" + ", ".join(f"{v:.3g}" for v in vals) + "]"


def fmt_float(value):
    if value is None:
        return "-"
    value = float(value)
    if abs(value) >= 1000 or (abs(value) < 1e-3 and value != 0):
        return f"{value:.3g}"
    if abs(value - round(value)) < 1e-10:
        return str(int(round(value)))
    return f"{value:.4g}"


def fmt_value(value):
    if isinstance(value, bool):
        return "True" if value else "False"
    if isinstance(value, str):
        return value
    return fmt_float(value)


def add_cover(doc):
    doc.add_paragraph()
    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rk = kicker.add_run("Technical Manual and User Guide")
    set_run_font(rk, size=11, bold=True, color=MUTED)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rt = title.add_run("ILQ Multi-Agent Intersection Simulator")
    set_run_font(rt, size=24, bold=True, color=BLUE)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = subtitle.add_run("Formulation, Solver Method, Current Configuration, and How-to-Use README")
    set_run_font(rs, size=13, italic=True, color=RGBColor(0x44, 0x44, 0x44))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = meta.add_run("Generated from repository source on June 5, 2026")
    set_run_font(rm, size=10.5, color=MUTED)

    doc.add_paragraph()
    add_callout(
        doc,
        "Purpose",
        "This document consolidates the simulator's mathematical formulation, active cost stack, ILQ solution method, line-search behavior, warm-start logic, output workflow, sensitivity-analysis tools, and practical usage instructions into one Word reference.",
    )
    doc.add_page_break()


def add_manual_contents(doc):
    doc.add_heading("Contents", level=1)
    for item in [
        "1. Introduction and simulator purpose",
        "2. Current scenario and configuration",
        "3. Mathematical formulation",
        "4. Cost functions and smoothing details",
        "5. Linearization, quadraticization, and LQ game solve",
        "6. ILQ iteration, line search, and convergence",
        "7. Warm start, arrival handling, and output generation",
        "8. How to use the simulator",
        "9. Modification guide and troubleshooting",
        "10. Appendices: live constants, agent weights, and source map",
    ]:
        add_bullet(doc, item)


def add_introduction(doc):
    doc.add_heading("1. Introduction", level=1)
    add_paragraph(
        doc,
        "The simulator models multiple autonomous or driver-like vehicles moving through a four-way intersection. Each vehicle is a strategic agent in a finite-horizon, nonzero-sum dynamic game. The agents share a concatenated state, but each player owns its own objective and chooses its own curvature and acceleration inputs.",
    )
    add_paragraph(
        doc,
        "The main experiment currently uses six vehicles: a southbound through vehicle, a northbound left-turn vehicle, their followers, and one right-turn vehicle on each approach. The scenario is solved with an iterative linear-quadratic game method using a feedback Nash LQ subproblem.",
    )
    add_callout(
        doc,
        "Design intent",
        "The model is not a microscopic traffic simulator. It is an optimal-control research simulator for testing how multi-agent objectives, warm starts, feedback Nash approximations, and smooth safety/yielding costs shape intersection behavior.",
    )
    doc.add_page_break()


def add_current_configuration(doc):
    doc.add_heading("2. Current Scenario and Configuration", level=1)
    core_constants = [
        ("Time step", main.DT, "seconds"),
        ("Horizon steps", main.HORIZON_STEPS, "samples"),
        ("Nominal horizon", main.HORIZON_STEPS * main.DT, "seconds"),
        ("State dimension per agent", main.STATE_DIM_PER_AGENT, "[px, py, speed, heading]"),
        ("Control dimension per agent", main.CONTROL_DIM_PER_AGENT, "[kappa, acceleration]"),
        ("Stop radius", main.STOP_RADIUS, "meters"),
        ("Pass-through radius", main.PASS_THROUGH_RADIUS, "meters"),
        ("Post-arrival output steps", main.STOP_AFTER_ARRIVAL_STEPS, "steps"),
        ("Static obstacle count", len(main.STATIC_OBSTACLES), "generated points"),
    ]
    solver_constants = [
        ("Max ILQ iterations", main.MAX_ITERATIONS, "iterations"),
        ("Convergence tolerance", main.CONVERGENCE_TOL, "max state update"),
        ("Base alpha scaling", main.ALPHA_SCALING, "feedforward scale"),
        ("Line search enabled", main.USE_ALPHA_LINE_SEARCH, "boolean"),
        ("Line-search minimum alpha", main.ALPHA_LINE_SEARCH_MIN, "alpha"),
        ("Line-search shrink", main.ALPHA_LINE_SEARCH_SHRINK, "multiplier"),
        ("Line-search max growth guard", main.ALPHA_LINE_SEARCH_MAX_GROWTH, "times previous max update"),
        ("Line-search start iteration", main.ALPHA_LINE_SEARCH_START_ITERATION, "iteration index"),
        ("LQ solver type", main.LQ_SOLVER_TYPE, "feedback or open_loop"),
    ]
    add_table(
        doc,
        ["Core item", "Current value", "Meaning"],
        [(name, fmt_value(value), unit) for name, value, unit in core_constants],
        [2.15, 1.55, 2.8],
    )
    add_table(
        doc,
        ["Solver item", "Current value", "Meaning"],
        [(name, fmt_value(value), unit) for name, value, unit in solver_constants],
        [2.15, 1.55, 2.8],
    )

    doc.add_heading("2.1 Road Geometry", level=2)
    geometry = [
        ("Intersection x range", f"{main.INTERSECTION_X_MIN} to {main.INTERSECTION_X_MAX}", "meters"),
        ("Intersection y range", f"{main.INTERSECTION_Y_MIN} to {main.INTERSECTION_Y_MAX}", "meters"),
        ("North-south divider x", main.NS_DIVIDER_X, "meters"),
        ("Southbound lane center x", main.SOUTHBOUND_LANE_X, "meters"),
        ("Northbound lane center x", main.NORTHBOUND_LANE_X, "meters"),
        ("East-west lane center y", main.EW_LANE_Y, "meters"),
        ("West arm x minimum", main.WEST_ARM_X_MIN, "meters"),
        ("East arm x maximum", main.EAST_ARM_X_MAX, "meters"),
        ("Road south y minimum", main.ROAD_SOUTH_Y_MIN, "meters"),
        ("Road north y maximum", main.ROAD_NORTH_Y_MAX, "meters"),
    ]
    add_table(doc, ["Geometry item", "Value", "Unit"], [(a, fmt_value(b), c) for a, b, c in geometry], [2.5, 1.8, 2.2])
    add_paragraph(
        doc,
        "Static obstacles are generated at 2 m spacing along curbs, approach boundaries, and the north-south divider, leaving a small divider gap near the intersection so turning vehicles can move through the box.",
    )

    doc.add_heading("2.2 Active Agents", level=2)
    rows = []
    for i, agent in enumerate(main.AGENTS):
        rows.append(
            (
                i,
                agent.name,
                fmt_array(agent.initial_state),
                fmt_array(agent.destination),
                fmt_float(agent.desired_speed),
            )
        )
    add_table(
        doc,
        ["#", "Agent", "Initial state", "Destination", "v_des"],
        rows,
        [0.35, 1.65, 2.0, 1.65, 0.85],
    )
    relation_rows = []
    for agent in main.AGENTS:
        if "through" in agent.name:
            role = "through"
        elif "right_turn" in agent.name:
            role = "right turn"
        elif "left_turn" in agent.name:
            role = "left turn"
        else:
            role = "follower"
        relation_rows.append(
            (
                agent.name,
                "-" if agent.leader_name is None else agent.leader_name,
                fmt_float(main.desired_terminal_heading(agent)),
                role,
            )
        )
    add_table(
        doc,
        ["Agent", "Leader", "Terminal heading", "Role"],
        relation_rows,
        [2.0, 2.05, 1.25, 1.2],
    )
    doc.add_page_break()


def add_formulation(doc):
    doc.add_heading("3. Mathematical Formulation", level=1)
    doc.add_heading("3.1 State, Control, and Horizon", level=2)
    add_paragraph(doc, "Each agent i has a four-dimensional state and a two-dimensional control:")
    add_formula(
        doc,
        "x_i = [px_i, py_i, v_i, theta_i]^T\n"
        "u_i = [kappa_i, a_i]^T\n"
        "x = [x_0^T, x_1^T, ..., x_{N-1}^T]^T in R^(4N)",
    )
    add_table(
        doc,
        ["Symbol", "Meaning", "Units"],
        [
            ("px_i, py_i", "Vehicle position in the road coordinate frame", "m"),
            ("v_i", "Longitudinal speed", "m/s"),
            ("theta_i", "Heading angle; 0 points along +x and pi/2 along +y", "rad"),
            ("kappa_i", "Path curvature / steering-like control", "rad/m"),
            ("a_i", "Longitudinal acceleration", "m/s^2"),
        ],
        [1.2, 4.15, 1.15],
    )

    doc.add_heading("3.2 Vehicle Dynamics", level=2)
    add_paragraph(
        doc,
        "Each vehicle uses a curvature-controlled kinematic model. The concatenated dynamics stack the independent vehicle dynamics; interaction enters through costs rather than direct physical coupling.",
    )
    add_formula(
        doc,
        "px_dot    = v cos(theta)\n"
        "py_dot    = v sin(theta)\n"
        "v_dot     = a\n"
        "theta_dot = v kappa\n\n"
        "x_dot = f(t, x, u_0, ..., u_{N-1})",
    )
    add_paragraph(
        doc,
        "Main solver runs pass use Euler integration for rollout and linearization consistency. The RK4 integrator remains available in ilq/rollout.py and can be selected by setting use_euler=False.",
    )

    doc.add_heading("3.3 Dynamic Game Objective", level=2)
    add_paragraph(
        doc,
        "Each player minimizes its own finite-horizon cost while responding to the other agents' controls and shared state trajectory. This is a nonzero-sum game, not a single global planner objective.",
    )
    add_formula(
        doc,
        "J_i = sum_{k=0}^{T-1} g_i(k, x_k, u_0,k, ..., u_{N-1,k}) + psi_i(x_T)\n\n"
        "x_{k+1} = F_d(k, x_k, u_0,k, ..., u_{N-1,k})",
    )

    doc.add_heading("3.4 Local Quadratic Game Approximation", level=2)
    add_paragraph(
        doc,
        "At each ILQ iteration, the nonlinear dynamics and costs are approximated around the current operating trajectory. The LQ solver uses local coordinates dx = x - x_ref and du_i = u_i - u_ref_i.",
    )
    add_formula(
        doc,
        "g_i approx 0.5 dx^T Q_i dx + l_i^T dx\n"
        "          + sum_j (0.5 du_j^T R_ij du_j + r_ij^T du_j)\n\n"
        "dx_{k+1} = A_k dx_k + sum_i B_i,k du_i,k",
    )
    add_paragraph(
        doc,
        "The QuadraticApprox container also has an S field for state-control cross terms, but pack_for_lq_game() currently does not pass S into either LQ solver. Cost terms therefore retain exact first-order gradients and use positive-semidefinite local curvature blocks for stability.",
    )


def add_costs(doc):
    doc.add_heading("4. Cost Functions and Smoothing Details", level=1)
    active_costs = [
        ("AgentRunningSpeedCost", "All agents, all steps", "Tracks desired speed and tapers target speed to zero near the destination."),
        ("AgentRunningDestinationCost", "All agents, all steps", "Weak running spring to the destination."),
        ("AgentArrivalHoldCost", "All agents, gated after arrival", "Keeps an arrived agent parked near its destination inside optimization."),
        ("BatchedStaticObstacleRepulsionCost", "All agents, all steps", "Repels each vehicle from generated curb/divider obstacles."),
        ("PairwiseAgentRepulsionCost", "Each ordered pair, weighted by leader/cross-traffic relationship", "Encodes distance-based collision/yield pressure with smooth moving gates."),
        ("AgentProximitySpeedCost", "Cross-traffic and direct-leader pairs when weights are positive", "Adds direct incentive to slow down when another vehicle is nearby."),
        ("AgentLeaderYieldLineCost", "Available for direct leader pairs when weight is positive", "Can hold a follower behind a line until its leader clears."),
        ("AgentLaneKeepingCost", "Agents with lane_x and positive lane weight", "Keeps through vehicles near lane center."),
        ("AgentSpeedDependentControlCost", "All agents, all steps", "Penalizes curvature with v^4 scaling and acceleration quadratically."),
        ("AgentFullTerminalCost", "Final horizon sample", "Anchors terminal position, speed, and heading."),
    ]
    add_table(doc, ["Cost term", "When used", "Role"], active_costs, [2.05, 2.05, 2.4])

    doc.add_heading("4.1 Core Running and Terminal Costs", level=2)
    add_formula(
        doc,
        "Running speed:\n"
        "  0.5 q_speed (v_i - v_target_i)^2\n\n"
        "Running destination:\n"
        "  0.5 w_dest ||p_i - p_goal_i||^2\n\n"
        "Arrival hold:\n"
        "  hold_gate * (0.5 w_pos ||p_i - p_goal_i||^2 + 0.5 w_speed v_i^2)\n\n"
        "Terminal:\n"
        "  0.5 b_px dx^2 + 0.5 b_py dy^2 + 0.5 b_speed (v - v_goal)^2\n"
        "  + 0.5 b_heading * (2 sin((theta - theta_goal)/2))^2",
    )
    add_paragraph(
        doc,
        "The speed target and arrival-hold gate use terminal-heading projected remaining distance, so the gate stays meaningful after a vehicle passes through the destination plane.",
    )

    doc.add_heading("4.2 Obstacle, Pairwise, and Yielding Costs", level=2)
    add_formula(
        doc,
        "Static obstacle repulsion:\n"
        "  sum_m w_obs / (||p_i - o_m||^2 + epsilon_obs)\n\n"
        "Pairwise repulsion:\n"
        "  w_pair gate_i gate_j / (||p_i - p_j||^2 + epsilon_pair)\n\n"
        "Proximity speed:\n"
        "  w_prox gate smooth_floor(v_i, 0, width) / (||p_i - p_j||^2 + epsilon_prox)\n\n"
        "Leader yield line:\n"
        "  gate(leader_clearance) * C2_hinge(follower_progress - hold_value)",
    )
    add_paragraph(
        doc,
        "Pairwise weights are semantic: leader_repulsion and leader_proximity_speed_weight apply only to a direct leader, while cross_traffic_repulsion and proximity_speed_weight apply to vehicles in different platoon roots.",
    )

    doc.add_heading("4.3 Control and Lane Costs", level=2)
    add_formula(
        doc,
        "Speed-dependent control:\n"
        "  0.5 beta_2 kappa_i^2 v_floor_i^4 + 0.5 beta_3 a_i^2\n\n"
        "Lane keeping:\n"
        "  0.5 w_lane (px_i - lane_x)^2  and/or  0.5 w_lane (py_i - lane_y)^2",
    )

    doc.add_page_break()
    doc.add_heading("4.4 Smooth Helper Functions", level=2)
    add_formula(
        doc,
        "C2 smootherstep:\n"
        "  S(s) = 0 for s <= 0, 1 for s >= 1,\n"
        "       = 6s^5 - 15s^4 + 10s^3 otherwise\n\n"
        "C2 speed floor transition:\n"
        "  p(s) = 3s^5 - 8s^4 + 6s^3\n\n"
        "Stable sigmoid:\n"
        "  sigmoid(z) = 1 / (1 + exp(-z)), evaluated with stable positive/negative branches\n\n"
        "C2 quadratic hinge:\n"
        "  0 when err <= 0, 0.5 w err^2 when err >= delta,\n"
        "  w delta^2 (1.5s^3 - 1.5s^4 + 0.5s^5) in the transition",
    )


def add_solution_method(doc):
    doc.add_heading("5. Solution Method", level=1)
    doc.add_heading("5.1 ILQ Iteration", level=2)
    add_ordered_list(doc, [
        "Roll out the current strategies from the fixed initial state x0.",
        "Evaluate each player's per-step costs and store the current operating point.",
        "Compute max state-trajectory change relative to the previous operating point.",
        "Linearize the discrete dynamics along the operating trajectory.",
        "Quadraticize each player's costs and pack local LQ terms.",
        "Solve the finite-horizon LQ dynamic game for feedback gains P_i and feedforward vectors alpha_i.",
        "Choose the feedforward scale with line search, then repeat until convergence or iteration cap.",
    ])

    doc.add_heading("5.2 Rollout Strategy", level=2)
    add_formula(
        doc,
        "u_i[k] = u_ref_i[k] - P_i[k] (x_k - x_ref[k]) - eta alpha_i[k]\n\n"
        "Safety clamps during rollout:\n"
        "  kappa in [-0.6, 0.6]\n"
        "  a     in [-8.0, 8.0]",
    )

    doc.add_heading("5.3 Linearization", level=2)
    add_formula(
        doc,
        "Central finite differences, eps = 1e-6:\n\n"
        "A[:,j] = (f(t, x + eps e_j, u) - f(t, x - eps e_j, u)) / (2 eps)\n\n"
        "B_i[:,j] = (f(t, x, ..., u_i + eps e_j, ...) - f(t, x, ..., u_i - eps e_j, ...)) / (2 eps)\n\n"
        "Euler discretization when use_euler=True:\n"
        "  A_d = I + dt A,  B_d,i = dt B_i",
    )
    add_paragraph(
        doc,
        "If use_euler=False, the simulator finite-differences the actual RK4 one-step map instead of discretizing the continuous Jacobian.",
    )

    doc.add_heading("5.4 Feedback Nash LQ Solver", level=2)
    add_paragraph(
        doc,
        "The current main experiment uses the feedback Nash LQ solver. It works backward from the terminal value functions and solves one coupled linear system per time step.",
    )
    add_formula(
        doc,
        "Local LQ dynamics:\n"
        "  dx+ = A dx + sum_i B_i du_i\n\n"
        "Value function per player:\n"
        "  V_i(dx) = 0.5 dx^T Z_i dx + zeta_i^T dx + const\n\n"
        "Coupled stationarity system:\n"
        "  S X = Y\n"
        "  first xdim columns of X -> P_i gains\n"
        "  final column of X -> alpha_i feedforward terms\n\n"
        "Returned strategy:\n"
        "  du_i = -P_i dx - alpha_i",
    )
    add_paragraph(
        doc,
        "The feedback solver applies Gershgorin-style diagonal regularization when the coupled Nash matrix is near-singular, increasing diagonal entries until each Gershgorin lower bound is at least 1e-3.",
    )

    doc.add_heading("5.5 Open-Loop Solver Availability", level=2)
    add_paragraph(
        doc,
        "The older open-loop LQ game solver is still present and can be selected with LQ_SOLVER_TYPE = \"open_loop\". It returns the same P_i and alpha_i structure but uses a different backward recursion. Current default is feedback.",
    )

    doc.add_heading("5.6 Line Search and Convergence", level=2)
    add_paragraph(
        doc,
        "Line search is applied to the feedforward scale eta. The current implementation accepts the largest finite candidate that reduces the realized summed player cost while satisfying a loose state-update runaway guard.",
    )
    add_formula(
        doc,
        "Candidate schedule:\n"
        "  eta = alpha_scaling, eta *= alpha_line_search_shrink until eta < alpha_line_search_min\n\n"
        "Merit:\n"
        "  merit(eta) = sum_i sum_k g_i(k, x_eta,k, u_eta,k)\n\n"
        "Accept if:\n"
        "  rollout is finite\n"
        "  max_delta_x <= alpha_line_search_max_growth * previous_max_delta, when previous_max_delta > tol\n"
        "  merit(eta) <= reference_merit - cost_tol * max(1, |reference_merit|)\n\n"
        "Convergence test:\n"
        "  max_k ||x_current[k] - x_previous[k]|| < convergence_tol",
    )
    add_callout(
        doc,
        "Important caveat",
        "The merit is a globalization heuristic for these traffic examples. A nonzero-sum game does not generally have a single true objective, so the summed cost is used pragmatically rather than interpreted as a potential-game proof.",
    )


def add_warm_start_outputs(doc):
    doc.add_heading("6. Warm Start, Arrival Handling, and Outputs", level=1)
    doc.add_heading("6.1 Nominal Warm Start", level=2)
    add_paragraph(
        doc,
        "make_initial_nominal_trajectory() creates a feasible destination-seeking rollout by integrating the same vehicle dynamics used by the solver. It does not hard-code inter-agent avoidance; pairwise behavior should come from the costs.",
    )
    add_formula(
        doc,
        "terminal_direction = [cos(theta_terminal), sin(theta_terminal)]\n"
        "remaining = (destination - position)^T terminal_direction\n"
        "lateral_error = ||delta - remaining terminal_direction||\n\n"
        "If lateral_error < pass_through_radius:\n"
        "  v_stop = sqrt(2 a_brake max(remaining - stop_buffer, 0))\n\n"
        "a = clip(2 (v_target - v), -4.0, 1.5)",
    )
    add_paragraph(
        doc,
        "Turning warm starts use approach waypoints to avoid cutting through lane dividers before entering the intersection, then use curvature proportional to heading or lateral error.",
    )

    doc.add_heading("6.2 Arrival Handling", level=2)
    add_paragraph(
        doc,
        "There are two terminal notions: the solver horizon terminal cost at the last sample, and per-agent pass-through arrival when the vehicle comes within the pass-through radius of its destination.",
    )
    add_bullet(doc, "Inside optimization: speed target tapers to zero, arrival hold activates, and pairwise gates deactivate near/past the terminal plane.")
    add_bullet(doc, "After optimization: freeze_arrived_agents() zeroes speed and controls after each agent's pass-through step for clean CSVs, plots, and GIFs.")

    doc.add_heading("6.3 Output Files", level=2)
    outputs = [
        ("trajectory_feedback.csv", "Step-by-step trajectory, controls, active flags, distance-to-destination, and pairwise distances."),
        ("trajectory_feedback.png", "Final road-scene trajectory plot."),
        ("trajectory_feedback_animation.gif", "Animated final trajectory."),
        ("feedback_costs_per_iteration.png", "Per-player cost convergence."),
        ("feedback_delta_x_per_iteration.png", "Max state-update convergence."),
        ("feedback_trajectory_evolution.png", "Selected trajectory evolution over ILQ iterations."),
        ("feedback_trajectory_per_agent.png", "Per-agent trajectory evolution across iterations."),
        ("sensitivity_analysis/...", "Batch-study outputs for setup, warm-start, and parameter sweeps."),
    ]
    add_paragraph(doc, "Main-run files are saved under outputs/ with the multi_agent_ prefix.")
    add_table(doc, ["Output suffix or folder", "Description"], outputs, [2.6, 3.9])


def add_how_to_use(doc):
    doc.add_heading("7. How to Use the Simulator", level=1)
    doc.add_heading("7.1 Quick Start", level=2)
    add_ordered_list(doc, [
        "Open a terminal at the repository root: /Users/amirmohammadkhakpour/Downloads/ILQR_Multi_Agent_Intersection",
        "Run the main scenario with: python main.py",
        "Review generated files under outputs/. The CSV is the most precise numerical artifact; PNG/GIF files are presentation and inspection artifacts.",
        "Run tests with: python -m pytest",
    ])

    add_formula(
        doc,
        "Common commands:\n"
        "  python main.py\n"
        "  python sensitivity_analysis.py --switch 1\n"
        "  python sensitivity_analysis.py --switch 2 --case-limit 2 --no-gifs\n"
        "  python sensitivity_analysis.py --switch 3 --max-iterations 100 --allow-early-convergence\n"
        "  python test_diagnosis_delay20.py --use-existing-csv\n"
        "  python -m pytest",
    )

    doc.add_heading("7.2 Sensitivity Switches", level=2)
    add_table(
        doc,
        ["Switch", "Study type", "Examples"],
        [
            ("1", "Different game setups", "Minimal two-agent case, four-agent followers, six-vehicle baseline, tight timing."),
            ("2", "Different initial guesses", "Nominal, zero-control, delayed turns, slow speed, perturbed reference."),
            ("3", "Different parameter values", "Interaction, control, obstacle, and speed-tracking weight scaling."),
        ],
        [0.65, 1.75, 4.1],
    )

    doc.add_heading("7.3 Reading the CSV", level=2)
    add_bullet(doc, "Each row is one time step until flexible game end: max first-arrival step plus STOP_AFTER_ARRIVAL_STEPS.")
    add_bullet(doc, "For each agent, the CSV records active flag, x, y, speed, heading, curvature, acceleration, and distance to destination.")
    add_bullet(doc, "Pairwise distance columns report center-to-center distances for each unordered agent pair.")

    doc.add_heading("7.4 Modifying the Scenario", level=2)
    add_ordered_list(doc, [
        "Edit the AGENTS list in main.py to add, remove, or change vehicles.",
        "Use leader_name instead of hard-coded leader_index when possible; leader references are resolved and validated by name.",
        "Tune each agent's CostWeights for speed tracking, control effort, obstacle repulsion, leader following, cross-traffic yielding, lane keeping, and terminal behavior.",
        "Change LQ_SOLVER_TYPE between \"feedback\" and \"open_loop\" to compare local game solvers.",
        "Adjust line-search parameters only after checking both cost convergence and trajectory quality.",
    ])

    doc.add_heading("7.5 Troubleshooting", level=2)
    add_table(
        doc,
        ["Symptom", "Likely cause", "What to inspect"],
        [
            ("Convergence flattens", "Small true updates, active constraints/gates, or merit fallback.", "alpha_history, line_search_history, cost plots, and per-agent trajectory evolution."),
            ("Vehicles overlap", "Pairwise/leader/proximity weights too weak or warm-start timing too tight.", "Minimum pairwise distances, diagnosis script, proximity-speed weights."),
            ("Vehicle clips road boundary", "Lane or obstacle weights too weak; warm start cuts a corner.", "Lane weights, obstacle map, trajectory evolution plot."),
            ("Controls saturate", "Feedback gains or cost Hessians demand aggressive correction.", "Max kappa/accel in summaries, control weights beta_2 and beta_3."),
            ("Long runtime", "Six-agent feedback solve plus finite differences and plotting/GIFs.", "Use --no-gifs, --case-limit, or lower max iterations for exploration."),
        ],
        [1.45, 2.25, 2.8],
    )


def add_appendices(doc):
    doc.add_page_break()
    doc.add_heading("8. Appendices", level=1)
    doc.add_heading("8.1 Cost Weights by Agent", level=2)
    rows = []
    for agent in main.AGENTS:
        w = agent.cost
        rows.append(
            (
                agent.name,
                fmt_float(w.q_speed),
                fmt_float(w.beta_2),
                fmt_float(w.beta_3),
                fmt_float(w.v_min),
                fmt_float(w.running_destination),
                fmt_float(w.static_obstacle_repulsion),
                fmt_float(w.static_obstacle_epsilon),
            )
        )
    add_table(
        doc,
        ["Agent", "q_speed", "beta_2", "beta_3", "v_min", "dest", "obs_w", "obs_eps"],
        rows,
        [1.55, 0.65, 0.6, 0.6, 0.55, 0.55, 0.75, 1.25],
    )

    rows = []
    for agent in main.AGENTS:
        w = agent.cost
        rows.append(
            (
                agent.name,
                fmt_float(w.leader_repulsion),
                fmt_float(w.leader_proximity_speed_weight),
                fmt_float(w.cross_traffic_repulsion),
                fmt_float(w.proximity_speed_weight),
                fmt_float(w.proximity_speed_epsilon),
                fmt_float(w.lane_x),
                fmt_float(w.lane_x_weight),
            )
        )
    add_table(
        doc,
        ["Agent", "lead_r", "lead_v", "cross_r", "prox_v", "prox_eps", "lane_x", "lane_w"],
        rows,
        [1.55, 0.72, 0.8, 0.7, 0.8, 0.62, 0.63, 0.68],
    )

    rows = []
    for agent in main.AGENTS:
        w = agent.cost
        rows.append(
            (
                agent.name,
                fmt_float(w.arrival_speed_transition),
                fmt_float(w.arrival_hold_transition),
                fmt_float(w.arrival_hold_position_weight),
                fmt_float(w.arrival_hold_speed_weight),
                fmt_float(w.terminal_position_weight),
                fmt_float(w.terminal_speed_weight),
                fmt_float(w.terminal_heading_weight),
            )
        )
    add_table(
        doc,
        ["Agent", "speed_tr", "hold_tr", "hold_p", "hold_v", "term_p", "term_v", "term_th"],
        rows,
        [1.55, 0.72, 0.72, 0.72, 0.72, 0.68, 0.68, 0.71],
    )

    doc.add_page_break()
    doc.add_heading("8.2 Source File Map", level=2)
    add_table(
        doc,
        ["File", "Role"],
        [
            ("main.py", "Scenario constants, agent definitions, game construction, warm start, plotting, CSV/GIF export."),
            ("ilq/ilq_solver.py", "Iterative LQ game loop, rollout, convergence, alpha line search."),
            ("ilq/linearization.py", "Finite-difference dynamics linearization and Euler/RK4 discrete-map linearization."),
            ("ilq/quadraticization.py", "Cost quadraticization packing into local LQ game terms."),
            ("lq_game/lq_feedback_solver.py", "Feedback Nash finite-horizon LQ game solver with Gershgorin regularization."),
            ("lq_game/lq_open_loop_solver.py", "Alternative open-loop/local LQ game solver."),
            ("costs/base_cost.py", "All scalar cost terms and local quadratic approximations."),
            ("costs/player_cost.py", "Player cost container with state/control regularizers."),
            ("dynamics/player_dynamics.py", "Vehicle dynamics implementations."),
            ("dynamics/base_dynamics.py", "Concatenated multi-agent dynamics wrapper."),
            ("sensitivity_analysis.py", "Batch sweeps over setup, initial guess, and parameter cases."),
            ("test_diagnosis_delay20.py", "Single-case delayed-turn diagnosis and overlap checks."),
            ("tests/", "Regression tests for config, derivatives, and line-search behavior."),
        ],
        [2.25, 4.25],
    )

    doc.add_heading("8.3 Known Modeling Assumptions", level=2)
    add_bullet(doc, "Vehicle interaction is cost-coupled, not physically collision-constrained.")
    add_bullet(doc, "State-control cross terms are not currently packed into the LQ game solvers.")
    add_bullet(doc, "Control clamps are applied during rollout as safety bounds.")
    add_bullet(doc, "The line-search merit is summed player cost, a practical heuristic rather than a game-theoretic potential proof.")
    add_bullet(doc, "Post-solve freezing is for clean reporting; the optimization also includes arrival behavior through costs.")


def build_doc():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_manual_contents(doc)
    add_introduction(doc)
    add_current_configuration(doc)
    add_formulation(doc)
    add_costs(doc)
    add_solution_method(doc)
    add_warm_start_outputs(doc)
    add_how_to_use(doc)
    add_appendices(doc)

    core_props = doc.core_properties
    core_props.title = "ILQ Multi-Agent Intersection Simulator Manual"
    core_props.subject = "Formulation, solution method, configuration, and user guide"
    core_props.author = "OpenAI Codex"
    core_props.comments = "Generated from repository source on June 5, 2026."
    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build_doc())
