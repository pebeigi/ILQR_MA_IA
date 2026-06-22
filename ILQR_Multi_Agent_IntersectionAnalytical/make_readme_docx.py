"""Generate README.docx — ILQ Game Optimal Control Formulation."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Style helpers ─────────────────────────────────────────────────────────────
def set_font(run, name="Calibri", size=11, bold=False, italic=False, color=None):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(text, level=1, color=(0x1F, 0x49, 0x7D)):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
        run.font.name = "Calibri"
    return p

def body(text, indent=0):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def math_block(text):
    """Monospace block for equations."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    set_font(run, name="Courier New", size=10, color=(0x20, 0x20, 0x20))
    # Light grey background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  "F2F2F2")
    pPr.append(shd)
    return p

def bullet(text, indent_cm=0.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Cm(indent_cm)
    run = p.add_run(text)
    set_font(run, size=11)
    return p

def add_table(headers, rows, col_widths=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1+len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_font(run, bold=True, size=10, color=(0xFF, 0xFF, 0xFF))
        # Blue fill
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "1F497D")
        tcPr.append(shd)
    # Data rows
    for r, row_data in enumerate(rows):
        cells = table.rows[r+1].cells
        fill = "DEEAF1" if r % 2 == 0 else "FFFFFF"
        for c, val in enumerate(row_data):
            cell = cells[c]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            set_font(run, size=10)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    return table


# ══════════════════════════════════════════════════════════════════════════════
#  TITLE
# ══════════════════════════════════════════════════════════════════════════════
p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("ILQ Multi-Agent Intersection Simulator")
set_font(r, size=20, bold=True, color=(0x1F, 0x49, 0x7D))

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Optimal Control & Game-Theoretic Formulation — Complete Reference")
set_font(r, size=12, italic=True, color=(0x40, 0x40, 0x40))

doc.add_paragraph()  # spacer

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
heading("1. Overview", level=1)
body(
    "This simulator models a multi-agent intersection scenario as a finite-horizon, "
    "N-player, nonzero-sum differential game solved via the Iterative Linear-Quadratic "
    "(ILQ) algorithm. Each vehicle agent independently minimizes its own cost functional "
    "while the dynamics of all agents are coupled through the shared state vector. The "
    "solver seeks an approximate open-loop Nash equilibrium by iteratively linearizing "
    "the dynamics and quadraticizing each player's cost, then solving the resulting "
    "Linear-Quadratic (LQ) game exactly via a backward Riccati recursion."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 2 — STATE & CONTROL
# ══════════════════════════════════════════════════════════════════════════════
heading("2. State Space and Control Space", level=1)

heading("2.1  Individual Agent State", level=2)
body("Each agent i has a 4-dimensional state vector:")
math_block("xᵢ = [ pxᵢ,  pyᵢ,  vᵢ,  θᵢ ]ᵀ  ∈  ℝ⁴")
add_table(
    ["Symbol", "Meaning", "Units"],
    [
        ["pxᵢ", "Longitudinal (east) position", "m"],
        ["pyᵢ", "Lateral (north) position", "m"],
        ["vᵢ",  "Longitudinal speed", "m/s"],
        ["θᵢ",  "Heading angle  (0 → +x axis, π/2 → +y axis)", "rad"],
    ],
    col_widths=[0.9, 3.8, 0.9]
)

heading("2.2  Concatenated State", level=2)
body("For N agents the full game state is the vertical concatenation:")
math_block("x = [ x₀ᵀ,  x₁ᵀ,  …,  x_{N-1}ᵀ ]ᵀ  ∈  ℝ^{4N}")

heading("2.3  Control Variables", level=2)
body("Each agent i has a 2-dimensional control input:")
math_block("uᵢ = [ κᵢ,  aᵢ ]ᵀ  ∈  ℝ²")
add_table(
    ["Symbol", "Meaning", "Units", "Clamp"],
    [
        ["κᵢ", "Path curvature  (= θ̇ᵢ / vᵢ)", "rad/m", "[-0.6, 0.6]"],
        ["aᵢ", "Longitudinal acceleration", "m/s²", "[-8.0, 8.0]"],
    ],
    col_widths=[0.7, 2.8, 0.8, 1.2]
)
body(
    "Controls are hard-clamped to these bounds inside the rollout loop to prevent "
    "runaway feedback gains from large proximity cost Hessians."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 3 — DYNAMICS
# ══════════════════════════════════════════════════════════════════════════════
heading("3. Continuous-Time Dynamics", level=1)
body(
    "Each agent obeys a kinematic bicycle model (curvature-controlled unicycle). "
    "The steering angle is implicitly encoded in the curvature κ, making the "
    "model independent of vehicle wheelbase."
)
math_block(
    "ẋᵢ = f( xᵢ, uᵢ ) =\n"
    "      ⎡  vᵢ · cos(θᵢ)  ⎤\n"
    "      ⎢  vᵢ · sin(θᵢ)  ⎥\n"
    "      ⎢       aᵢ        ⎥\n"
    "      ⎣   vᵢ · κᵢ      ⎦"
)
body(
    "The concatenated dynamics F(x, u₀, …, u_{N-1}) stacks the N individual "
    "dynamics functions independently (no direct mechanical coupling between agents; "
    "coupling enters only through the cost functions)."
)

heading("3.1  Discrete-Time Integration", level=2)
add_table(
    ["Parameter", "Value", "Description"],
    [
        ["dt", "0.1 s",   "Time step"],
        ["T",  "200",     "Horizon steps  (20 s total)"],
        ["Integrator", "RK4 (default)", "4th-order Runge-Kutta; Euler available"],
    ],
    col_widths=[1.2, 1.4, 3.6]
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 4 — GAME STRUCTURE
# ══════════════════════════════════════════════════════════════════════════════
heading("4. Game Structure", level=1)
body(
    "The problem is formulated as a finite-horizon, N-player, nonzero-sum "
    "differential game. Each player i independently minimizes:"
)
math_block(
    "min_{uᵢ₀, …, uᵢ_{T-1}}   Jᵢ  =  Σₖ₌₀^{T-1} ℓᵢ(xₖ, {uₖ}, k)  +  φᵢ(x_T)"
)
body(
    "subject to the concatenated nonlinear dynamics. Players do not coordinate "
    "— the solution concept is an approximate open-loop Nash equilibrium, found "
    "by the ILQ solver. Each player's cost depends on the full state x (which "
    "contains all agents' positions and speeds), so the games are non-trivially coupled."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 5 — RUNNING COST
# ══════════════════════════════════════════════════════════════════════════════
heading("5. Running Cost  ℓᵢ(x, u, k)", level=1)
body(
    "The running cost is a sum of terms. Each term has an associated weight; "
    "setting a weight to zero disables that term. Not every term is active for "
    "every agent — see Section 7 for agent-specific configurations."
)

# 5a
heading("5a.  Destination Spring  (AgentRunningDestinationCost)", level=2)
body("Pulls the agent toward its goal position at every step:")
math_block("ℓ_dest = (β₁/2) · ‖ [pxᵢ, pyᵢ] − dᵢ ‖²")
add_table(
    ["Symbol", "Meaning", "Default"],
    [["β₁", "running_destination weight", "1.0"], ["dᵢ", "2D goal position [dx, dy]", "agent-specific"]],
    col_widths=[0.8, 3.2, 1.5]
)

# 5b
heading("5b.  Speed Tracking  (AgentRunningSpeedCost)", level=2)
body("Penalizes deviation from a desired cruising speed:")
math_block("ℓ_speed = (q_speed/2) · (vᵢ − v*)²")
add_table(
    ["Symbol", "Meaning", "Through / Follower", "Left-Turn"],
    [
        ["q_speed", "Speed tracking weight", "200.0", "20.0"],
        ["v*", "Desired speed", "6.0 m/s", "6.0 m/s"],
    ],
    col_widths=[1.0, 2.4, 1.5, 1.2]
)
body(
    "A lower q_speed for the left-turn vehicle allows it to genuinely slow down "
    "when yielding to through traffic."
)

# 5c
heading("5c.  State-Dependent Control Cost  (AgentSpeedDependentControlCost)", level=2)
body(
    "Penalizes both curvature and acceleration, with the curvature penalty "
    "scaling as v⁴ to represent lateral-force costs at high speed:"
)
math_block("ℓ_ctrl = (β₂/2) · κᵢ² · vᵢ⁴  +  (β₃/2) · aᵢ²")
body("In the local quadratic approximation this becomes a state-dependent control weight matrix:")
math_block("R(vᵢ) = diag( β₂ vᵢ⁴ + ε ,  β₃ + ε )   (frozen at nominal vᵢ each iteration)")
add_table(
    ["Symbol", "Through / Follower", "Left-Turn", "Description"],
    [
        ["β₂", "1.5",  "10.0", "Curvature cost coefficient"],
        ["β₃", "80.0",  "8.0", "Acceleration cost coefficient"],
        ["v_min", "6.0 m/s", "1.0 m/s", "Speed floor to avoid v⁴=0 singularity"],
        ["ε",  "0.001",  "0.001", "Regularization added to R diagonal"],
    ],
    col_widths=[1.0, 1.4, 1.2, 2.6]
)
body(
    "High β₂ for the left-turn vehicle is deliberately avoided — it was found "
    "empirically that high β₂ makes the vehicle swing east (rightward) before "
    "turning, worsening the conflict."
)

# 5d
heading("5d.  Static Obstacle Repulsion  (BatchedStaticObstacleRepulsionCost)", level=2)
body("Repels the agent from road-boundary cell obstacles:")
math_block(
    "ℓ_obs = Σ_{obs ∈ Obstacles}  w_obs / ( ‖[pxᵢ,pyᵢ] − obs‖²  +  ε_obs )"
)
body(
    "All obstacle points are processed in a single vectorized pass. "
    "The LQ approximation uses a positive-semidefinite (PSD) curvature matrix "
    "to ensure the subproblem remains well-posed near obstacles."
)
add_table(
    ["Parameter", "Value"],
    [["w_obs", "1000.0"], ["ε_obs", "0.1"]],
    col_widths=[1.5, 1.5]
)

# 5e
heading("5e.  Pairwise Agent Repulsion  (PairwiseAgentRepulsionCost)", level=2)
body("Repels agent i away from agent j:")
math_block("ℓ_rep(i,j) = w_rep / ( ‖pᵢ − pⱼ‖²  +  ε_rep )")
body("Three semantically distinct weight tiers are used:")
bullet("leader_repulsion — gap-keeping: follower → its direct leader in the same platoon")
bullet("cross_traffic_repulsion — yielding: agents in different platoon streams")
bullet("other_agent_repulsion — catch-all for remaining pairs")
body(
    "The cost is deactivated when either agent is within stop_radius = 2.0 m of "
    "its destination, preventing repulsion from fighting against arrival."
)
body("PSD approximation of the Hessian in relative-position coordinates:")
math_block(
    "Q_block = curvature · [ [ I₂, -I₂ ],      curvature = max(ε, 2·w / z²)\n"
    "                         [ -I₂,  I₂ ] ]    z = ‖pᵢ−pⱼ‖² + ε_rep"
)

# 5f
heading("5f.  Proximity Speed Cost  (AgentProximitySpeedCost) — Yielding Mechanism", level=2)
body(
    "This is the primary yielding term. It couples the speed of agent i with "
    "proximity to agent j, creating a direct gradient to slow down:"
)
math_block("ℓ_yield = w_prox · max(vᵢ, 0) / ( ‖pᵢ − pⱼ‖²  +  ε_prox )")
body(
    "When pᵢ is close to pⱼ, the coefficient on vᵢ rises sharply. "
    "This produces yielding behavior that PairwiseAgentRepulsionCost alone "
    "cannot achieve because repulsion only penalizes position, not speed."
)
add_table(
    ["Parameter", "Left-Turn Vehicle"],
    [["w_prox", "3000"], ["ε_prox", "2.0"]],
    col_widths=[1.5, 1.5]
)
body("Active only for cross-traffic pairs (left-turn vs. through + follower).")

# 5g
heading("5g.  Lane Keeping  (AgentLaneKeepingCost)", level=2)
body("Soft quadratic constraint on the lane centre:")
math_block("ℓ_lane = (w_lane/2) · (pxᵢ − x_lane)²")
add_table(
    ["Agent", "x_lane (m)", "w_lane"],
    [
        ["through_southbound",  "4.5", "200.0"],
        ["follower_southbound", "4.5", "200.0"],
        ["left_turn_northbound","—",   "disabled"],
    ],
    col_widths=[2.5, 1.2, 1.2]
)

# 5h
heading("5h.  Soft Geometric Barriers  (Left-Turn Vehicle Only)", level=2)
body("One-sided quadratic penalties that enforce road geometry constraints:")
math_block(
    "Y-ceiling:  ℓ_yc = (w_y/2) · max(0,  pyᵢ − y_max)²     [blocks northward overshoot]\n"
    "X-ceiling:  ℓ_xc = (w_x/2) · max(0,  pxᵢ − x_max)²     [blocks rightward swing before turn]"
)
add_table(
    ["Barrier", "Bound", "Weight", "Purpose"],
    [
        ["Y-ceiling", "y_max = 16.0 m", "500.0",  "Limits post-turn northward drift (WB lane top ≈ 16 m)"],
        ["X-ceiling", "x_max = 13.5 m", "5000.0", "Prevents east swing before turn (NB lane centre)"],
    ],
    col_widths=[1.2, 1.5, 0.9, 2.6]
)
body(
    "These are zero below the bound and quadratic above — they act as "
    "soft one-sided walls without requiring hard constraints."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 6 — TERMINAL COST
# ══════════════════════════════════════════════════════════════════════════════
heading("6. Terminal Cost  φᵢ(x_T)  — Active Only at Step k = T", level=1)
body("A four-component terminal cost penalizes all state dimensions at the final step:")
math_block(
    "φᵢ =  (b_px/2) · (pxᵢ − px*)²\n"
    "    + (b_py/2) · (pyᵢ − py*)²\n"
    "    + (b_speed/2) · (vᵢ − v*)²\n"
    "    + (b_heading/2) · wrap(θᵢ − θ*)²"
)
body(
    "The heading error is wrapped to [−π, π] via arctan2 before squaring, "
    "ensuring numerical stability across heading discontinuities. In the "
    "quadratic approximation, the q vector absorbs the wrapped error offset."
)
add_table(
    ["Weight", "Default Value", "Description"],
    [
        ["b_px",      "250.0", "Terminal x-position error weight"],
        ["b_py",      "250.0", "Terminal y-position error weight"],
        ["b_speed",     "5.0", "Terminal speed error weight"],
        ["b_heading",  "50.0", "Terminal heading error weight"],
        ["px*, py*", "destination", "Target final position"],
        ["v*",   "desired_speed", "Target final speed"],
        ["θ*",   "desired_heading", "Target final heading"],
    ],
    col_widths=[1.3, 1.5, 3.4]
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 7 — AGENT CONFIGURATIONS
# ══════════════════════════════════════════════════════════════════════════════
heading("7. Agent Configurations (Current Scenario)", level=1)
body(
    "The intersection geometry is: N-S road x ∈ [0, 18] m with centreline at x = 9 m. "
    "SB lane centre x = 4.5 m, NB lane centre x = 13.5 m. "
    "E-W road y ∈ [9, 21] m, WB lane centre y = 15 m."
)
doc.add_paragraph()

heading("Agent 0 — through_southbound", level=2)
add_table(
    ["Parameter", "Value"],
    [
        ["Initial state", "[4.5, 38.0, 6.0, −π/2]  (x, y, v, θ)"],
        ["Destination",   "[4.5, −28.0]"],
        ["Desired speed", "6.0 m/s"],
        ["q_speed",       "200.0"],
        ["β₂, β₃",        "1.5, 80.0"],
        ["Lane keeping",  "x_lane = 4.5 m,  w_lane = 200"],
        ["Repulsion",     "All disabled — has right-of-way, ignores others"],
        ["Role",          "Southbound through vehicle, platoon root"],
    ],
    col_widths=[2.0, 4.2]
)

doc.add_paragraph()
heading("Agent 1 — left_turn_northbound", level=2)
add_table(
    ["Parameter", "Value"],
    [
        ["Initial state", "[13.5, −7.0, 6.0, π/2]  (x, y, v, θ)"],
        ["Destination",   "[−16.0, 15.0]  (west, in WB lane)"],
        ["Desired speed", "6.0 m/s"],
        ["q_speed",       "20.0  (weak — allows genuine slowing)"],
        ["β₂, β₃",        "10.0, 8.0  (low β₂ prevents rightward swing)"],
        ["cross_traffic_repulsion", "20.0  (yields to SB stream)"],
        ["proximity_speed_weight", "3000  (primary yield mechanism)"],
        ["Y-ceiling",     "y_max = 16.0 m,  weight = 500"],
        ["X-ceiling",     "x_max = 13.5 m,  weight = 5000"],
        ["Role",          "NB→WB left-turn, yields to through traffic"],
    ],
    col_widths=[2.5, 3.7]
)

doc.add_paragraph()
heading("Agent 2 — follower_southbound", level=2)
add_table(
    ["Parameter", "Value"],
    [
        ["Initial state", "[4.5, 46.0, 6.0, −π/2]  (x, y, v, θ)"],
        ["Destination",   "[4.5, −20.0]  (8 m further south than leader)"],
        ["Desired speed", "6.0 m/s"],
        ["q_speed",       "200.0"],
        ["β₂, β₃",        "1.5, 80.0"],
        ["leader_repulsion", "60.0  (gap-keeping behind Agent 0)"],
        ["Lane keeping",  "x_lane = 4.5 m,  w_lane = 200"],
        ["Role",          "SB platoon follower, gap-keeps behind through vehicle"],
    ],
    col_widths=[2.0, 4.2]
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 8 — QUADRATIC APPROXIMATION
# ══════════════════════════════════════════════════════════════════════════════
heading("8. Local Quadratic Cost Approximation", level=1)
body(
    "At each ILQR iteration, every cost term is quadraticized around the "
    "current nominal trajectory. The local approximation for player i's cost "
    "at step k takes the form:"
)
math_block(
    "ℓ̂ᵢ(δx, {δuⱼ}) ≈  (1/2) δxᵀ Qᵢ δx  +  lᵢᵀ δx\n"
    "                +  Σⱼ  [ (1/2) δuⱼᵀ Rᵢⱼ δuⱼ  +  rᵢⱼᵀ δuⱼ ]\n"
    "                +  const"
)
body("where δx = x − x̄ and δuⱼ = uⱼ − ūⱼ are deviations from the nominal trajectory.")
add_table(
    ["Symbol", "Dimension", "Meaning"],
    [
        ["Qᵢ",  "4N × 4N",  "State cost Hessian for player i"],
        ["lᵢ",  "4N",       "State cost gradient for player i"],
        ["Rᵢⱼ", "2 × 2",    "Control cost Hessian (player i's cost, player j's control)"],
        ["rᵢⱼ", "2",        "Control cost gradient"],
    ],
    col_widths=[0.9, 1.4, 3.9]
)
body(
    "All Hessian matrices are forced to be positive semidefinite (PSD). "
    "Non-convex terms (obstacle and agent repulsion) use a PSD curvature "
    "lower-bounded by 1×10⁻⁸ to ensure the LQ subproblem remains solvable."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 9 — ILQ SOLVER
# ══════════════════════════════════════════════════════════════════════════════
heading("9. ILQ Solver Algorithm", level=1)

heading("9.1  Outer Loop", level=2)
body("Repeat until convergence or max iterations:")
bullet("Step 1 — Rollout: apply current feedback policy to produce nominal {x̄ₖ, ūᵢₖ}")
bullet("Step 2 — Linearize dynamics at each step: xₖ₊₁ ≈ Aₖ xₖ + Σᵢ Bᵢₖ uᵢₖ")
bullet("Step 3 — Quadraticize costs: compute Qᵢₖ, lᵢₖ, Rᵢⱼₖ, rᵢⱼₖ for all i, k")
bullet("Step 4 — Backward Riccati pass: solve the LQ game for feedback gains Pᵢₖ and feedforward αᵢₖ")
bullet("Step 5 — Update policy and check convergence")

heading("9.2  Feedback Policy", level=2)
body("The policy applied during the rollout is:")
math_block(
    "uᵢₖ  =  ūᵢₖ  −  Pᵢₖ (xₖ − x̄ₖ)  −  α_scale · αᵢₖ\n\n"
    "α_scale = 0.05   (step-size damping for stability)"
)

heading("9.3  Backward Riccati Pass (LQ Game)", level=2)
body("At each step k (backwards from T to 0), form the joint optimality system:")
math_block(
    "S_k · [ P₀ₖᵀ ; P₁ₖᵀ ; … ; P_{N-1,k}ᵀ ]  =  Y_k\n\n"
    "S_k is block matrix of size (Σᵢ uᵢ_dim) × (Σᵢ uᵢ_dim):\n"
    "  S[i,i]  =  Rᵢᵢ  +  Bᵢᵀ Zᵢ Bᵢ\n"
    "  S[i,j]  =  Bᵢᵀ Zᵢ Bⱼ ,  i ≠ j\n\n"
    "Y_k  =  [ B₀ᵀ Z₀ A ; B₁ᵀ Z₁ A ; … ]"
)
body("Value function update (closed-loop matrix F = A − Σᵢ Bᵢ Pᵢ):")
math_block(
    "Zᵢ ← Fᵀ Zᵢ F  +  Qᵢ  +  Σⱼ Pⱼᵀ Rᵢⱼ Pⱼ\n"
    "ζᵢ ← Fᵀ (ζᵢ + Zᵢ β)  +  lᵢ  +  Σⱼ Pⱼᵀ Rᵢⱼ αⱼ  −  Σⱼ Pⱼᵀ rᵢⱼ\n\n"
    "β  =  −Σᵢ Bᵢ αᵢ   (closed-loop affine term)"
)
body("The system S is solved via least-squares (numpy.linalg.lstsq) for numerical robustness.")

heading("9.4  Convergence Criterion", level=2)
math_block(
    "max_k  ‖x̄ₖ^{new} − x̄ₖ^{old}‖₂  <  δ_tol\n\n"
    "δ_tol = 1×10⁻⁴,   max iterations = 25"
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 10 — NOTATION SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading("10. Notation Summary", level=1)
add_table(
    ["Symbol", "Meaning"],
    [
        ["N",           "Number of agents (players)"],
        ["T",           "Horizon steps (200 steps = 20 s)"],
        ["dt",          "Time step (0.1 s)"],
        ["xᵢ ∈ ℝ⁴",    "State of agent i: [px, py, v, θ]"],
        ["x ∈ ℝ^{4N}", "Concatenated game state"],
        ["uᵢ ∈ ℝ²",    "Control of agent i: [κ, a]"],
        ["κ",           "Path curvature (rad/m), θ̇/v"],
        ["a",           "Longitudinal acceleration (m/s²)"],
        ["dᵢ",          "2D destination of agent i"],
        ["v*",          "Desired cruise speed"],
        ["θ*",          "Desired terminal heading"],
        ["Jᵢ",          "Total cost for player i"],
        ["ℓᵢ",          "Running cost for player i at step k"],
        ["φᵢ",          "Terminal cost for player i at step T"],
        ["Aₖ",          "Linearized dynamics Jacobian ∂f/∂x at step k"],
        ["Bᵢₖ",         "Control Jacobian ∂f/∂uᵢ at step k"],
        ["Qᵢₖ",         "State cost Hessian (quadraticized) for player i at step k"],
        ["Rᵢⱼₖ",        "Control cost Hessian for player i w.r.t. player j's control"],
        ["Pᵢₖ",         "Feedback gain matrix for player i at step k"],
        ["αᵢₖ",         "Feedforward control correction for player i at step k"],
        ["Zᵢₖ",         "Value function Hessian for player i at step k"],
        ["ζᵢₖ",         "Value function gradient (linear term) for player i at step k"],
        ["β₁",          "Running destination spring weight"],
        ["β₂",          "Curvature cost coefficient in R(v)"],
        ["β₃",          "Acceleration cost coefficient in R(v)"],
        ["w_obs",        "Static obstacle repulsion weight"],
        ["w_rep",        "Pairwise agent repulsion weight"],
        ["w_prox",       "Proximity speed (yielding) weight"],
        ["ε",           "Regularization / softening epsilon in barrier costs"],
    ],
    col_widths=[1.8, 4.4]
)

# ══════════════════════════════════════════════════════════════════════════════
#  SECTION 11 — KEY DESIGN CHOICES
# ══════════════════════════════════════════════════════════════════════════════
heading("11. Key Design Choices and Assumptions", level=1)
bullet(
    "Kinematic model: each agent is a kinematic bicycle (no inertia, no tire slip). "
    "This is valid at low-to-moderate urban speeds."
)
bullet(
    "Nonzero-sum game: agents are not cooperative — the through vehicle has right-of-way "
    "and ignores others; the left-turn vehicle must infer and yield on its own."
)
bullet(
    "PSD Hessian forcing: all non-convex cost Hessians (repulsion, barriers) are "
    "lower-bounded to ensure the backward pass LQ subproblem is always solvable."
)
bullet(
    "State-dependent R(v): the curvature cost β₂κ²v⁴ is frozen at the nominal speed "
    "each outer iteration. This is a first-order approximation that is exact only at the "
    "linearization point."
)
bullet(
    "Heading wrap: the terminal heading error is wrapped to [−π, π] via arctan2 at each "
    "quadraticization call, handling the π ↔ −π discontinuity."
)
bullet(
    "Yielding mechanism: ProximitySpeedCost (w·v/dist²) is the primary yield signal. "
    "PairwiseRepulsionCost alone cannot produce speed reduction because it acts only "
    "on position gradients."
)
bullet(
    "Open-loop Nash equilibrium: the ILQ solver finds a fixed point of the iterative "
    "Riccati recursion. This is a local Nash equilibrium — global optimality is not "
    "guaranteed for nonlinear, non-convex games."
)
bullet(
    "Control clamping: κ ∈ [−0.6, 0.6] rad/m and a ∈ [−8, 8] m/s² are enforced in "
    "every rollout step. Large proximity cost Hessians can otherwise generate gain matrices "
    "that produce unrealizable accelerations."
)

# ══════════════════════════════════════════════════════════════════════════════
#  SAVE
# ══════════════════════════════════════════════════════════════════════════════
out_path = "README_Formulation.docx"
doc.save(out_path)
print(f"Saved: {out_path}")